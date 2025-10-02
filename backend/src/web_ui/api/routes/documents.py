"""
Document management API routes for React frontend.

Provides CRUD operations for documents with ChromaDB integration
and DocumentEditingAgent support.
"""

from typing import Any
import asyncio
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from ...agent.document_editor import DocumentEditingAgent
from ...utils.logging_config import get_logger
from ..dependencies import get_document_agent

logger = get_logger(__name__)

# Create router
router = APIRouter()

# --- Enhanced Pydantic Models for Document API Requests/Responses ---


class DocumentCreateRequest(BaseModel):
    filename: str
    content: str = ""
    document_type: str = "markdown"
    metadata: dict[str, Any] | None = None


class DocumentType(BaseModel):
    """Enhanced document type definitions"""

    type: str  # markdown, richtext, plaintext, code, json, pdf
    language: str | None = None  # For code files
    editor_mode: str = "source"  # source, visual, split


class DocumentCreateLiveRequest(BaseModel):
    title: str
    content: str = ""
    document_type: str = "markdown"
    metadata: dict[str, Any] | None = None


class DocumentEditLiveRequest(BaseModel):
    title: str | None = None
    content: str
    metadata: dict[str, Any] | None = None


class DocumentExportRequest(BaseModel):
    document_id: str
    content: str
    format: str  # pdf, html, docx
    document_type: str = "markdown"


class DocumentEditRequest(BaseModel):
    document_id: str
    instruction: str
    use_llm: bool = True


class DocumentSearchRequest(BaseModel):
    query: str
    collection_type: str = "documents"
    limit: int = 10
    use_mcp_tools: bool = True


class DocumentSuggestionsRequest(BaseModel):
    content: str
    document_type: str = "document"


class PolicyStoreRequest(BaseModel):
    document_id: str
    policy_title: str
    policy_type: str = "manual"
    authority_level: str = "medium"


class BatchProcessRequest(BaseModel):
    file_paths: list[str]
    document_type: str = "document"


class ChatRequest(BaseModel):
    """Request model for chat messages."""

    message: str
    context_document_id: str | None = None


class ChatResponse(BaseModel):
    """Response model for chat messages."""

    response: str


class DocumentVersionRequest(BaseModel):
    document_id: str
    version: int


# --- Enhanced Document Management Endpoints ---


@router.post("/create-live")
async def create_document_live(
    request: DocumentCreateLiveRequest,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Create a new document with live editing support."""
    try:
        # Generate unique filename if not provided
        filename = f"{request.title}.{get_file_extension(request.document_type)}"

        success, message, document_id = await agent.create_document(
            filename=filename,
            content=request.content,
            document_type=request.document_type,
            metadata={
                **(request.metadata or {}),
                "title": request.title,
                "created_at": datetime.now().isoformat(),
                "document_type": request.document_type,
                "version": 1,
            },
        )

        if success:
            document = agent.chroma_manager.get_document("documents", document_id)
            return {
                "success": True,
                "id": document_id,
                "title": request.title,
                "content": request.content,
                "document_type": request.document_type,
                "created_at": datetime.now().isoformat(),
                "updated_at": datetime.now().isoformat(),
                "owner_id": "user",  # Replace with actual user ID from auth
                "metadata": document.metadata if document else {},
            }
        else:
            raise HTTPException(status_code=400, detail=message)
    except Exception as e:
        logger.error(f"Error creating live document: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error creating document: {str(e)}")


@router.put("/edit-live/{document_id}")
async def edit_document_live(
    document_id: str,
    request: DocumentEditLiveRequest,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Edit a document with live editing support."""
    try:
        # Get existing document
        existing_doc = agent.chroma_manager.get_document("documents", document_id)
        if not existing_doc:
            raise HTTPException(status_code=404, detail="Document not found")

        # Update metadata
        updated_metadata = {
            **existing_doc.metadata,
            **(request.metadata or {}),
            "updated_at": datetime.now().isoformat(),
            "version": existing_doc.metadata.get("version", 1) + 1,
        }

        if request.title:
            updated_metadata["title"] = request.title

        # Use agent to update document
        success, message, _ = await agent.edit_document(
            document_id=document_id,
            instruction=f"Update content to: {request.content}",
            use_llm=False,  # Direct content update, no LLM processing
        )

        if success:
            # Update metadata separately
            agent.chroma_manager.update_document_metadata(
                "documents", document_id, updated_metadata
            )

            return {
                "success": True,
                "message": "Document updated successfully",
                "document_id": document_id,
                "updated_at": updated_metadata["updated_at"],
                "version": updated_metadata["version"],
            }
        else:
            raise HTTPException(status_code=400, detail=message)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error editing live document {document_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error updating document: {str(e)}")


@router.post("/export")
async def export_document(
    request: DocumentExportRequest,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Export document to various formats (PDF, HTML, DOCX)."""
    try:
        # Import export libraries
        import io
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        import markdown
        from docx import Document as DocxDocument

        content = request.content
        document_type = request.document_type
        export_format = request.format.lower()

        if export_format == "pdf":
            # Create PDF
            buffer = io.BytesIO()
            p = canvas.Canvas(buffer, pagesize=letter)

            # Simple text-to-PDF conversion
            y_position = 750
            lines = content.split("\n")

            for line in lines:
                if y_position < 50:  # Start new page
                    p.showPage()
                    y_position = 750

                # Handle long lines
                if len(line) > 80:
                    words = line.split(" ")
                    current_line = ""
                    for word in words:
                        if len(current_line + word) < 80:
                            current_line += word + " "
                        else:
                            p.drawString(72, y_position, current_line.strip())
                            y_position -= 20
                            current_line = word + " "
                            if y_position < 50:
                                p.showPage()
                                y_position = 750
                    if current_line:
                        p.drawString(72, y_position, current_line.strip())
                        y_position -= 20
                else:
                    p.drawString(72, y_position, line)
                    y_position -= 20

            p.save()
            buffer.seek(0)

            return StreamingResponse(
                io.BytesIO(buffer.getvalue()),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=document.pdf"},
            )

        elif export_format == "html":
            # Convert to HTML
            if document_type == "markdown":
                html_content = markdown.markdown(content)
            else:
                html_content = f"<pre>{content}</pre>"

            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>Document Export</title>
                <style>
                    body {{ font-family: Arial, sans-serif; line-height: 1.6; padding: 20px; }}
                    pre {{ background: #f4f4f4; padding: 10px; border-radius: 5px; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """

            return StreamingResponse(
                io.BytesIO(full_html.encode()),
                media_type="text/html",
                headers={"Content-Disposition": f"attachment; filename=document.html"},
            )

        elif export_format == "docx":
            # Create DOCX
            doc = DocxDocument()

            # Add content paragraphs
            for line in content.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph()

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return StreamingResponse(
                io.BytesIO(buffer.getvalue()),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=document.docx"},
            )

        else:
            raise HTTPException(
                status_code=400, detail=f"Unsupported export format: {export_format}"
            )

    except Exception as e:
        logger.error(f"Error exporting document: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error exporting document: {str(e)}")


@router.get("/types")
async def get_supported_document_types():
    """Get list of supported document types and their configurations."""
    return {
        "document_types": [
            {
                "type": "markdown",
                "extensions": [".md", ".markdown"],
                "language": "markdown",
                "supports_preview": True,
                "editor_modes": ["source", "visual", "split"],
            },
            {
                "type": "richtext",
                "extensions": [".html", ".htm"],
                "language": "html",
                "supports_preview": True,
                "editor_modes": ["source", "visual"],
            },
            {
                "type": "plaintext",
                "extensions": [".txt"],
                "language": "plaintext",
                "supports_preview": False,
                "editor_modes": ["source"],
            },
            {
                "type": "code",
                "extensions": [".js", ".ts", ".py", ".css", ".json", ".yaml", ".xml"],
                "language": "auto-detect",
                "supports_preview": False,
                "editor_modes": ["source"],
            },
            {
                "type": "json",
                "extensions": [".json"],
                "language": "json",
                "supports_preview": False,
                "editor_modes": ["source"],
            },
        ],
        "export_formats": ["pdf", "html", "docx"],
        "languages": [
            "javascript",
            "typescript",
            "python",
            "html",
            "css",
            "json",
            "yaml",
            "xml",
            "markdown",
            "plaintext",
        ],
    }


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    document_type: str = "auto-detect",
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Upload a document file."""
    try:
        # Read file content
        content = await file.read()

        # Detect document type if auto-detect
        if document_type == "auto-detect":
            document_type = detect_document_type_from_filename(file.filename or "")

        # Decode content based on file type
        try:
            text_content = content.decode("utf-8")
        except UnicodeDecodeError:
            # Handle binary files
            if file.filename and file.filename.endswith(".pdf"):
                text_content = "[PDF file - binary content]"
                document_type = "pdf"
            else:
                raise HTTPException(status_code=400, detail="Unsupported file format")

        # Create document
        success, message, document_id = await agent.create_document(
            filename=file.filename or "uploaded_file",
            content=text_content,
            document_type=document_type,
            metadata={
                "uploaded_at": datetime.now().isoformat(),
                "original_filename": file.filename,
                "file_size": len(content),
                "content_type": file.content_type,
            },
        )

        if success:
            return {
                "success": True,
                "document_id": document_id,
                "filename": file.filename,
                "document_type": document_type,
                "size": len(content),
            }
        else:
            raise HTTPException(status_code=400, detail=message)

    except Exception as e:
        logger.error(f"Error uploading document: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error uploading document: {str(e)}")


@router.get("/{document_id}/versions")
async def get_document_versions(
    document_id: str,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Get version history of a document."""
    try:
        # This would require implementing version tracking in ChromaDB
        # For now, return current version info
        document = agent.chroma_manager.get_document("documents", document_id)
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        return {
            "document_id": document_id,
            "current_version": document.metadata.get("version", 1),
            "versions": [
                {
                    "version": document.metadata.get("version", 1),
                    "created_at": document.metadata.get(
                        "updated_at", document.timestamp.isoformat()
                    ),
                    "size": len(document.content),
                    "is_current": True,
                }
            ],
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting document versions: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error getting versions: {str(e)}")


# --- Document Management Endpoints ---


@router.post("/create")
async def create_document(
    request: DocumentCreateRequest,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Create a new document using the DocumentEditingAgent."""
    try:
        success, message, document_id = await agent.create_document(
            filename=request.filename,
            content=request.content,
            document_type=request.document_type,
            metadata=request.metadata,
        )
        if success:
            return {"success": True, "message": message, "document_id": document_id}
        else:
            raise HTTPException(status_code=400, detail=message)
    except Exception as e:
        logger.error(f"Error creating document: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, detail=f"Error creating document: {str(e)}"
        )


@router.post("/edit")
async def edit_document(
    request: DocumentEditRequest,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Edit a document using the DocumentEditingAgent."""
    try:
        success, message, document_id = await agent.edit_document(
            document_id=request.document_id,
            instruction=request.instruction,
            use_llm=request.use_llm,
        )
        if success:
            document = agent.chroma_manager.get_document(
                "documents", document_id or request.document_id
            )
            return {
                "success": True,
                "message": message,
                "document_id": document_id,
                "content": document.content if document else None,
                "metadata": document.metadata if document else None,
            }
        else:
            raise HTTPException(status_code=400, detail=message)
    except Exception as e:
        logger.error(f"Error editing document: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error editing document: {str(e)}")


@router.post("/search")
async def search_documents(
    request: DocumentSearchRequest,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Search documents using the DocumentEditingAgent."""
    try:
        results = await agent.search_documents(
            query=request.query,
            collection_type=request.collection_type,
            limit=request.limit,
            use_mcp_tools=request.use_mcp_tools,
        )
        return {"success": True, "results": results, "total": len(results)}
    except Exception as e:
        logger.error(f"Error searching documents: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, detail=f"Error searching documents: {str(e)}"
        )


@router.post("/suggestions")
async def get_document_suggestions(
    request: DocumentSuggestionsRequest,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Get document suggestions using the DocumentEditingAgent."""
    try:
        suggestions = await agent.get_document_suggestions(
            content=request.content, document_type=request.document_type
        )
        return {"success": True, "suggestions": suggestions}
    except Exception as e:
        logger.error(f"Error getting suggestions: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, detail=f"Error getting suggestions: {str(e)}"
        )


@router.post("/batch")
async def process_batch_documents(
    request: BatchProcessRequest,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Process multiple documents in a batch."""
    try:
        results = await agent.process_batch_documents(
            file_paths=request.file_paths, document_type=request.document_type
        )
        return {"success": True, "results": results}
    except Exception as e:
        logger.error(f"Error in batch processing: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, detail=f"Error in batch processing: {str(e)}"
        )


@router.post("/store-policy")
async def store_as_policy(
    request: PolicyStoreRequest,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Store a document as a policy manual."""
    try:
        success, message = await agent.store_as_policy(
            document_id=request.document_id,
            policy_title=request.policy_title,
            policy_type=request.policy_type,
            authority_level=request.authority_level,
        )
        if success:
            return {"success": True, "message": message}
        else:
            raise HTTPException(status_code=400, detail=message)
    except Exception as e:
        logger.error(f"Error storing policy: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error storing policy: {str(e)}")


@router.get("/{document_id}")
async def get_document(
    document_id: str,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Get a specific document by its ID."""
    try:
        document = agent.chroma_manager.get_document("documents", document_id)
        if not document:
            raise HTTPException(
                status_code=404, detail=f"Document not found: {document_id}"
            )
        return {
            "id": document.id,
            "content": document.content,
            "metadata": document.metadata,
            "created_at": document.timestamp.isoformat(),
            "updated_at": document.metadata.get(
                "updated_at", document.timestamp.isoformat()
            ),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting document {document_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error getting document: {str(e)}")


@router.delete("/{document_id}")
async def delete_document(
    document_id: str,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Delete a document by its ID."""
    try:
        success = agent.chroma_manager.delete_document("documents", document_id)
        if success:
            return {"success": True, "message": f"Document {document_id} deleted"}
        else:
            raise HTTPException(
                status_code=404, detail=f"Document not found: {document_id}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document {document_id}: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, detail=f"Error deleting document: {str(e)}"
        )


@router.get("/")
async def list_documents(
    collection_type: str = "documents",
    limit: int = 100,
    offset: int = 0,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """List all documents in a collection with pagination."""
    try:
        documents = agent.chroma_manager.get_all_documents(collection_type)
        total = len(documents)
        paginated_docs = documents[offset : offset + limit]
        return {
            "documents": [
                {
                    "id": doc.id,
                    "name": doc.metadata.get(
                        "filename", doc.metadata.get("title", "Untitled")
                    ),
                    "content": doc.content,
                    "metadata": doc.metadata,
                    "created_at": doc.timestamp.isoformat(),
                    "updated_at": doc.metadata.get(
                        "updated_at", doc.timestamp.isoformat()
                    ),
                    "type": collection_type,
                }
                for doc in paginated_docs
            ],
            "total": total,
            "limit": limit,
            "offset": offset,
        }
    except Exception as e:
        logger.error(f"Error listing documents: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, detail=f"Error listing documents: {str(e)}"
        )


@router.post("/chat", response_model=ChatResponse)
async def chat_with_document_agent(
    request: ChatRequest,
    agent: DocumentEditingAgent = Depends(get_document_agent),
):
    """Chat with the document agent."""
    try:
        response = await agent.chat_with_user(
            message=request.message,
            context_document_id=request.context_document_id,
        )
        return ChatResponse(response=response)
    except Exception as e:
        logger.error(f"Error in document chat endpoint: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to get chat response")


def get_file_extension(document_type: str) -> str:
    """Get file extension for document type."""
    extensions = {
        "markdown": "md",
        "richtext": "html",
        "plaintext": "txt",
        "code": "txt",
        "json": "json",
        "pdf": "pdf",
    }
    return extensions.get(document_type, "txt")


def detect_document_type_from_filename(filename: str) -> str:
    """Detect document type from filename."""
    if not filename:
        return "plaintext"

    ext = filename.split(".")[-1].lower()
    type_map = {
        "md": "markdown",
        "markdown": "markdown",
        "html": "richtext",
        "htm": "richtext",
        "txt": "plaintext",
        "js": "code",
        "ts": "code",
        "py": "code",
        "css": "code",
        "json": "json",
        "yaml": "code",
        "yml": "code",
        "xml": "code",
        "pdf": "pdf",
    }
    return type_map.get(ext, "plaintext")
